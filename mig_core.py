from pathlib import Path
from io import BytesIO

import pandas as pd
from docx import Document

PROMPT_TEMPLATE_DIR = Path("prompt_templates")
TEMPLATE_DIR = Path("sablony")


def make_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in str(text or "").splitlines():
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_mig_card_row(
    name: str,
    code: str,
    ean: str,
    price: float,
    standard_price: float,
    product_type: str,
    description: str = "",
) -> pd.DataFrame:
    price = float(price or 0)
    standard_price = float(standard_price or 0)
    price_without_vat = round(price / 1.21, 2) if price else 0

    row = {
        "code": str(code or "").strip(),
        "pairCode": "",
        "externalCode": str(code or "").strip(),
        "name": str(name or "").strip(),
        "ean": str(ean or "").strip(),
        "price": price,
        "priceWithoutVat": price_without_vat,
        "standardPrice": standard_price,
        "description": str(description or "").strip(),
        "manufacturer": "AMMO by MIG",
        "availabilityInStock": "Skladem",
        "availabilityOutOfStock": "Na dotaz",
        "googleCategoryIdInFeed": 6000,
        "heurekaCategoryId": 2351,
        "zboziCategoryId": 2413,
        "googleCategoryId": 6000,
        "itemType": str(product_type or "").strip(),
        "filteringProperty:Použití": "",
        "filteringProperty:Tvar štětce": "",
        "filteringProperty:Typ štětin": "",
        "filteringProperty:Velikost": "",
    }

    return pd.DataFrame([row])


def build_mig_prompt(
    prompt_type: str,
    product_name: str,
    product_ean: str,
    product_code: str = "",
    lang: str = "cs",
) -> str:
    lang = str(lang or "cs").strip().lower()
    template_path = PROMPT_TEMPLATE_DIR / f"{prompt_type}_{lang}.txt"

    if not template_path.exists():
        raise FileNotFoundError(f"Šablona nenalezena: {template_path}")

    template_text = template_path.read_text(encoding="utf-8")

    return f"""{template_text}

--------------------------------------------------
PRODUKT
{product_name or ""}

EAN
{product_ean or ""}

CODE
{product_code or ""}
--------------------------------------------------
"""


def parse_key_value_block(block_text: str) -> dict[str, str]:
    result: dict[str, str] = {}
    current_key = None
    current_value_lines: list[str] = []

    for raw_line in str(block_text or "").splitlines():
        line = raw_line.rstrip()

        if ":" in line:
            maybe_key, maybe_value = line.split(":", 1)
            key = maybe_key.strip()

            if key:
                if current_key is not None:
                    result[current_key] = "\n".join(current_value_lines).strip()

                current_key = key
                current_value_lines = [maybe_value.strip()]
                continue

        if current_key is not None:
            current_value_lines.append(line.strip())

    if current_key is not None:
        result[current_key] = "\n".join(current_value_lines).strip()

    return result


def replace_placeholders_in_docx(template_path: Path, values: dict[str, str]) -> str:
    if not template_path.exists():
        raise FileNotFoundError(f"Šablona DOCX nenalezena: {template_path}")

    doc = Document(template_path)
    normalized_values = {str(k): str(v or "") for k, v in values.items()}

    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text

        for key, value in normalized_values.items():
            new_text = new_text.replace("{" + key + "}", value)

        if new_text != full_text:
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            paragraph.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    new_text = full_text

                    for key, value in normalized_values.items():
                        new_text = new_text.replace("{" + key + "}", value)

                    if new_text != full_text:
                        for i in range(len(paragraph.runs) - 1, -1, -1):
                            paragraph._element.remove(paragraph.runs[i]._element)
                        paragraph.add_run(new_text)

    html_lines: list[str] = []

    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt:
            html_lines.append(txt)

    for table in doc.tables:
        html_lines.append("<table>")
        for row in table.rows:
            html_lines.append("<tr>")
            for cell in row.cells:
                html_lines.append(f"<td>{cell.text.strip()}</td>")
            html_lines.append("</tr>")
        html_lines.append("</table>")

    return "\n".join(html_lines)


def get_mig_template_paths(template_kind: str) -> tuple[dict[str, Path], dict[str, Path]]:
    if template_kind == "mig_paints":
        short_files = {
            "cs": TEMPLATE_DIR / "BARVA kratky popis.docx",
            "en": TEMPLATE_DIR / "BARVA kratky popis en.docx",
            "sk": TEMPLATE_DIR / "BARVA kratky popis sk.docx",
        }
        long_files = {
            "cs": TEMPLATE_DIR / "BARVA dlouhy popis.docx",
            "en": TEMPLATE_DIR / "BARVA dlouhy popis en.docx",
            "sk": TEMPLATE_DIR / "BARVA dlouhy popis sk.docx",
        }
    elif template_kind == "mig_tools":
        short_files = {
            "cs": TEMPLATE_DIR / "stetce - kratky text.docx",
            "en": TEMPLATE_DIR / "stetce - kratky text en.docx",
            "sk": TEMPLATE_DIR / "stetce - kratky text sk.docx",
        }
        long_files = {
            "cs": TEMPLATE_DIR / "stetce - dlouhy text.docx",
            "en": TEMPLATE_DIR / "stetce - dlouhy text en.docx",
            "sk": TEMPLATE_DIR / "stetce - dlouhy text sk.docx",
        }
    else:
        raise ValueError(f"Neznámý template_kind: {template_kind}")

    return short_files, long_files


def build_mig_html_single_lang(
    ai_output: str,
    template_kind: str,
    lang: str,
    extra_values: dict[str, str] | None = None,
) -> dict[str, str]:
    lang = str(lang or "cs").strip().lower()
    short_files, long_files = get_mig_template_paths(template_kind)

    if lang not in short_files or lang not in long_files:
        raise ValueError(f"Neplatný jazyk: {lang}")

    values = parse_key_value_block(ai_output)

    if extra_values:
        values.update({str(k): str(v or "") for k, v in extra_values.items()})

    out: dict[str, str] = {}

    out[f"shortDescription:{lang}"] = replace_placeholders_in_docx(short_files[lang], values)
    out[f"description:{lang}"] = replace_placeholders_in_docx(long_files[lang], values)

    product_name = str(values.get("nazev_produktu", "")).strip()
    short_desc = str(values.get("strucny_popis_produktu", "")).strip()

    if product_name and short_desc.lower().startswith(product_name.lower()):
        short_desc = short_desc[len(product_name):].strip(" -–—,:;")

    final_short = f"{product_name} {short_desc}".strip()
    final_short = final_short.replace("\n", " ").replace("\r", " ")

    out[f"name:{lang}"] = product_name
    out[f"seoTitle:{lang}"] = f"{product_name} | AMMO by MIG" if product_name else ""
    out[f"xmlFeedName:{lang}"] = product_name
    out[f"metaDescription:{lang}"] = final_short[:155] if final_short else ""

    return out


def apply_mig_output_to_csv(
    df: pd.DataFrame,
    row_index: int,
    ai_output: str,
    template_kind: str,
    lang: str,
    extra_values: dict[str, str] | None = None,
) -> pd.DataFrame:
    html_map = build_mig_html_single_lang(
        ai_output=ai_output,
        template_kind=template_kind,
        lang=lang,
        extra_values=extra_values,
    )
    df_out = df.copy()

    for col, value in html_map.items():
        if col not in df_out.columns:
            df_out[col] = ""
        df_out.at[row_index, col] = value

    if extra_values:
        if "img1_src" in extra_values:
            if "image" not in df_out.columns:
                df_out["image"] = ""
            df_out.at[row_index, "image"] = extra_values["img1_src"]

        if "img2_src" in extra_values:
            if "image2" not in df_out.columns:
                df_out["image2"] = ""
            df_out.at[row_index, "image2"] = extra_values["img2_src"]

        if "img3_src" in extra_values:
            if "image3" not in df_out.columns:
                df_out["image3"] = ""
            df_out.at[row_index, "image3"] = extra_values["img3_src"]

    df_out = df_out.drop(
        columns=[c for c in ["name", "description", "shortDescription"] if c in df_out.columns],
        errors="ignore",
    )

    preferred_order = [
        "code",
        "pairCode",
        "externalCode",

        "name:cs",
        "name:en",
        "name:sk",

        "shortDescription:cs",
        "shortDescription:en",
        "shortDescription:sk",

        "description:cs",
        "description:en",
        "description:sk",

        "price",
        "priceWithoutVat",
        "standardPrice",

        "categoryText",
        "warranty",
        "supplier",
        "manufacturer",
        "itemType",

        "filteringProperty:Použití",
        "filteringProperty:Tvar štětce",
        "filteringProperty:Typ štětin",
        "filteringProperty:Velikost",

        "googleCategoryIdInFeed",
        "heurekaCategoryId",
        "zboziCategoryId",
        "googleCategoryId",

        "image",
        "image2",
        "image3",

        "stock",
        "percentVat",
        "ossTaxRate:CZ",

        "availabilityInStock",
        "availabilityOutOfStock",

        "ean",
        "productVisibility",

        "xmlFeedName:cs",
        "xmlFeedName:en",
        "xmlFeedName:sk",

        "seoTitle:cs",
        "seoTitle:en",
        "seoTitle:sk",

        "metaDescription:cs",
        "metaDescription:en",
        "metaDescription:sk",
    ]

    existing_preferred = [col for col in preferred_order if col in df_out.columns]
    remaining_cols = [col for col in df_out.columns if col not in existing_preferred]

    df_out = df_out[existing_preferred + remaining_cols]
    df_out = df_out.fillna("")

    return df_out